#!/usr/bin/env python3
"""
生成 OpenRoboOLP 使用说明的 Word 文档
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from datetime import datetime

def create_document():
    # 创建文档
    doc = Document()
    
    # 设置文档属性
    doc.core_properties.title = "OpenRoboOLP 使用说明"
    doc.core_properties.author = "OpenRoboOLP Team"
    doc.core_properties.subject = "机器人离线编程软件使用手册"
    doc.core_properties.created = datetime.now()
    doc.core_properties.version = "v0.1.0"
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10.5)
    
    # 添加标题页
    add_title_page(doc)
    
    # 添加目录
    add_table_of_contents(doc)
    
    # 添加章节
    add_section_1(doc)   # 软件介绍
    add_section_2(doc)   # 系统要求
    add_section_3(doc)   # 快速开始
    add_section_4(doc)   # 界面说明
    add_section_5(doc)   # 功能特性
    add_section_6(doc)   # 常见问题
    add_section_7(doc)   # 技术支持
    
    return doc

def add_title_page(doc):
    """添加标题页"""
    # 添加新节（标题页）
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("OpenRoboOLP")
    title_run.font.name = '微软雅黑'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run("机器人离线编程软件")
    subtitle_run.font.name = '微软雅黑'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    subtitle_run.font.size = Pt(16)
    
    # 版本信息
    version = doc.add_paragraph()
    version.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    version_run = version.add_run("版本 v0.1.0")
    version_run.font.name = '微软雅黑'
    version_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    version_run.font.size = Pt(12)
    
    # 空行
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 版权信息
    copyright = doc.add_paragraph()
    copyright.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    copyright_run = copyright.add_run("© 2026 OpenRoboOLP Team. All rights reserved.")
    copyright_run.font.name = '微软雅黑'
    copyright_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    copyright_run.font.size = Pt(10)
    
    # 添加新节（正文）
    doc.add_section(WD_SECTION.NEW_PAGE)

def add_table_of_contents(doc):
    """添加目录"""
    heading = doc.add_heading("目录", level=1)
    heading_run = heading.runs[0]
    heading_run.font.name = '微软雅黑'
    heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 添加目录项
    toc_items = [
        ("软件介绍", 2, "1"),
        ("系统要求", 2, "2"),
        ("快速开始", 2, "3"),
        ("界面说明", 2, "4"),
        ("功能特性", 2, "5"),
        ("常见问题", 2, "6"),
        ("技术支持", 2, "7"),
    ]
    
    for item, level, number in toc_items:
        p = doc.add_paragraph(style='List Paragraph')
        p.add_run(f"{number}. {item}").font.name = '微软雅黑'
        p.add_run('\t' * 5).font.name = '微软雅黑'
        p.add_run('..').font.name = '微软雅黑'
        p.add_run(' 1').font.name = '微软雅黑'
    
    doc.add_section(WD_SECTION.NEW_PAGE)

def add_section_1(doc):
    """第1章：软件介绍"""
    heading = doc.add_heading("1. 软件介绍", level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    p1 = doc.add_paragraph()
    p1.add_run("OpenRoboOLP 是一款基于 Python 和 PySide6 开发的机器人离线编程软件，")
    p1.add_run("提供直观的 3D 可视化界面，支持机器人模型加载、路径规划、")
    p1.add_run("仿真演示等功能，帮助用户快速完成机器人离线编程任务。")
    
    p2 = doc.add_paragraph()
    p2.add_run("主要功能包括：")
    
    features = [
        "3D 机器人模型可视化",
        "关节角度实时调节",
        "路径录制与播放",
        "路径代码编辑",
        "路径保存与加载",
        "机器人属性查看",
    ]
    
    for feature in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(feature).font.name = '微软雅黑'

def add_section_2(doc):
    """第2章：系统要求"""
    heading = doc.add_heading("2. 系统要求", level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 创建表格
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    requirements = [
        ("操作系统", "Windows 10 或更高版本（64位）"),
        ("处理器", "支持 SSE2 指令集的处理器"),
        ("内存", "建议 2 GB 或更多"),
        ("磁盘空间", "至少 200 MB 可用空间"),
    ]
    
    for i, (key, value) in enumerate(requirements):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        
        # 设置字体
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_section_3(doc):
    """第3章：快速开始"""
    heading = doc.add_heading("3. 快速开始", level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 3.1 运行软件
    sub_heading = doc.add_heading("3.1 运行软件", level=2)
    sub_heading.runs[0].font.name = '微软雅黑'
    
    p = doc.add_paragraph()
    p.add_run("双击 ")
    p.add_run("OpenRoboOLP.exe").bold = True
    p.add_run(" 文件即可直接运行，无需安装任何依赖。")
    
    # 3.2 加载机器人模型
    sub_heading = doc.add_heading("3.2 加载机器人模型", level=2)
    sub_heading.runs[0].font.name = '微软雅黑'
    
    steps = [
        "点击菜单 ",
        "File",
        " → ",
        "Open Model",
        "",
        "选择 URDF 格式的机器人模型文件",
        "项目中已包含示例模型在 ",
        "models/",
        " 目录下",
    ]
    
    for i, step in enumerate(steps):
        p = doc.add_paragraph(style='List Number')
        if step.startswith('File') or step.startswith('Open') or step.startswith('models'):
            run = p.add_run(step)
            run.bold = True
        else:
            p.add_run(step)
    
    # 3.3 路径编程与仿真
    sub_heading = doc.add_heading("3.3 路径编程与仿真", level=2)
    sub_heading.runs[0].font.name = '微软雅黑'
    
    # 录制路径
    p = doc.add_paragraph()
    p.add_run("**录制路径**：").bold = True
    p.add_run("在右侧属性面板调整机器人关节角度，点击左侧 Record 按钮录制当前位置。")
    
    # 播放路径
    p = doc.add_paragraph()
    p.add_run("**播放路径**：").bold = True
    p.add_run("点击 Play 按钮开始播放录制的路径动画。")
    
    # 路径代码格式
    p = doc.add_paragraph()
    p.add_run("**路径代码格式**：").bold = True
    code = doc.add_paragraph("movej([0, -90, 90, -90, 90, 0], vel=0.5)")
    code.style = 'No Spacing'
    
    # 3.4 保存/加载路径
    sub_heading = doc.add_heading("3.4 保存/加载路径", level=2)
    sub_heading.runs[0].font.name = '微软雅黑'
    
    p = doc.add_paragraph()
    p.add_run("点击菜单 ").bold = False
    p.add_run("File").bold = True
    p.add_run(" → ").bold = False
    p.add_run("Save Path Code").bold = True
    p.add_run(" 保存路径代码到文件。").bold = False
    
    p = doc.add_paragraph()
    p.add_run("点击菜单 ").bold = False
    p.add_run("File").bold = True
    p.add_run(" → ").bold = False
    p.add_run("Load Path Code").bold = True
    p.add_run(" 加载已保存的路径代码。").bold = False

def add_section_4(doc):
    """第4章：界面说明"""
    heading = doc.add_heading("4. 界面说明", level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 界面布局描述
    p = doc.add_paragraph()
    p.add_run("软件界面分为三个主要区域：")
    
    areas = [
        ("左侧：路径代码编辑器", "用于编辑和显示路径代码，包含录制、播放、停止、清除按钮"),
        ("中间：3D 机器人显示区", "显示机器人模型的 3D 视图，支持旋转和缩放"),
        ("右侧：属性编辑区", "显示机器人关节角度和属性信息，支持实时调节"),
    ]
    
    for area, desc in areas:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"**{area}**：").bold = True
        p.add_run(desc)
    
    # 快捷键表格
    p = doc.add_paragraph()
    p.add_run("**快捷键**：")
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    shortcuts = [
        ("Ctrl+O", "加载路径代码"),
        ("Ctrl+S", "保存路径代码"),
        ("R", "重置 3D 视图"),
        ("Ctrl+Q", "退出程序"),
    ]
    
    for i, (key, value) in enumerate(shortcuts):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_section_5(doc):
    """第5章：功能特性"""
    heading = doc.add_heading("5. 功能特性", level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    features = [
        ("3D 机器人可视化", [
            "支持 URDF 格式机器人模型",
            "实时关节角度更新",
            "可旋转、缩放的 3D 视图",
        ]),
        ("路径录制与播放", [
            "多位置录制",
            "路径可视化显示",
            "可调播放速度",
        ]),
        ("路径代码编辑", [
            "支持手动编辑路径代码",
            "实时解析代码更新路径",
            "路径代码可保存和加载",
        ]),
        ("关节属性编辑", [
            "滑块调节关节角度",
            "数值输入精确控制",
            "显示当前关节状态",
        ]),
    ]
    
    for feature_name, items in features:
        sub_heading = doc.add_heading(f"5.{features.index((feature_name, items)) + 1} {feature_name}", level=2)
        sub_heading.runs[0].font.name = '微软雅黑'
        
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item).font.name = '微软雅黑'

def add_section_6(doc):
    """第6章：常见问题"""
    heading = doc.add_heading("6. 常见问题", level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    faqs = [
        ("软件启动后没有反应？", "请确保系统满足最低要求，并尝试以管理员权限运行。"),
        ("3D 视图不显示机器人？", "请确保已加载 URDF 模型文件，检查文件路径是否正确。"),
        ("录制的路径播放不正确？", "请确保所有记录的位置都在机器人的运动范围内。"),
        ("如何自定义机器人模型？", "使用 URDF 格式创建机器人模型，然后通过 File → Open Model 加载。"),
    ]
    
    for question, answer in faqs:
        p = doc.add_paragraph()
        p.add_run(f"**Q：{question}**").bold = True
        
        p = doc.add_paragraph()
        p.add_run(f"A：{answer}")

def add_section_7(doc):
    """第7章：技术支持"""
    heading = doc.add_heading("7. 技术支持", level=1)
    heading.runs[0].font.name = '微软雅黑'
    heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    p = doc.add_paragraph()
    p.add_run("如需技术支持，请查阅项目 README.md 文件或访问项目仓库。")
    
    p = doc.add_paragraph()
    p.add_run("**祝您使用愉快！**").bold = True

if __name__ == "__main__":
    doc = create_document()
    output_path = "dist/OpenRoboOLP使用说明.docx"
    doc.save(output_path)
    print(f"Word 文档已生成：{output_path}")
